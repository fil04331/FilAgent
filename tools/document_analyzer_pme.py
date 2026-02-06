"""
Outil d'analyse de documents pour PME québécoises
Supporte: PDF, Excel, Word avec calculs TPS/TVQ
"""

import pandas as pd
from typing import Dict, Any, Optional
import PyPDF2
import docx
import re
import logging
from pathlib import Path
from .base import BaseTool, ToolResult, ToolStatus

# Setup standard Python logger for testing compatibility
logger = logging.getLogger(__name__)

# Import custom logger for compliance (Loi 25 - audit trail)
try:
    from runtime.middleware.logging import get_logger
except ImportError:
    get_logger = None


class DocumentAnalyzerPME(BaseTool):
    """Analyseur intelligent de documents PME avec conformité Loi 25"""

    def __init__(self):
        super().__init__(
            name="document_analyzer_pme",
            description="Analyse de documents PME avec calculs TPS/TVQ (Québec)",
        )
        self.tps_rate = 0.05  # 5%
        self.tvq_rate = 0.09975  # 9.975%

        # Initialize logger for compliance (audit trail)
        try:
            self.logger = get_logger() if get_logger else None
        except Exception:
            self.logger = None

        # Security: Allowed paths for document analysis (similar to FileReaderTool)
        self.allowed_paths = [
            "working_set/",
            "temp/",
            "memory/working_set/",
            "/tmp/",  # Gradio uploads to system temp
        ]
        self.max_file_size = 50 * 1024 * 1024  # 50 MB (as per documentation)

    def _redact_file_path(self, file_path: str) -> str:
        """
        Redact PII from file paths for compliance (Loi 25, PIPEDA)

        Removes:
        - Usernames (e.g., /Users/john.doe/ → /Users/[REDACTED]/)
        - Sensitive folder names (confidential, secret, private, etc.)
        - Email addresses in filenames
        - Phone numbers in filenames
        - SSN patterns in filenames

        Args:
            file_path: Original file path potentially containing PII

        Returns:
            Redacted file path safe for logging
        """
        path = Path(file_path)

        # Sensitive folder/file name patterns
        sensitive_patterns = [
            "confidential",
            "secret",
            "private",
            "personal",
            "prive",
            "confidentiel",
        ]

        # Redact parts in path
        parts = list(path.parts)
        for i, part in enumerate(parts):
            # Redact username in path (e.g., /Users/john.doe/)
            if i > 0 and parts[i - 1].lower() in ["users", "home"]:
                parts[i] = "[REDACTED]"
            # Redact sensitive folder names
            elif any(sensitive in part.lower() for sensitive in sensitive_patterns):
                parts[i] = "[REDACTED]"

        # Redact filename if it contains PII patterns
        filename = path.name

        # Email pattern
        filename = re.sub(
            r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}", "[EMAIL_REDACTED]", filename
        )

        # Phone pattern (514-555-1234, 5145551234)
        filename = re.sub(r"\d{3}[-.\s]?\d{3}[-.\s]?\d{4}", "[PHONE_REDACTED]", filename)

        # SSN pattern (123-45-6789)
        filename = re.sub(r"\d{3}-\d{2}-\d{4}", "[SSN_REDACTED]", filename)

        # Sensitive words in filename (with underscores/hyphens)
        for sensitive in sensitive_patterns:
            # Match word with boundaries OR with underscore/hyphen
            filename = re.sub(
                rf"(?:^|[^a-zA-Z]){sensitive}(?:[^a-zA-Z]|$)",
                "[REDACTED]",
                filename,
                flags=re.IGNORECASE,
            )

        # Reconstruct path
        parts[-1] = filename
        return str(Path(*parts))

    def _redact_pii_from_error(self, error_message: str) -> str:
        """
        Redact PII from error messages for compliance

        Args:
            error_message: Original error message

        Returns:
            Redacted error message safe for logging
        """
        # Email redaction
        error_message = re.sub(
            r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}", "[EMAIL_REDACTED]", error_message
        )

        # Phone redaction
        error_message = re.sub(r"\d{3}[-.\s]?\d{3}[-.\s]?\d{4}", "[PHONE_REDACTED]", error_message)

        # SSN redaction
        error_message = re.sub(r"\d{3}-\d{2}-\d{4}", "[SSN_REDACTED]", error_message)

        # Path redaction (any absolute path)
        error_message = re.sub(r"/(?:Users|home)/[^/\s]+", "/[REDACTED]", error_message)

        return error_message

    def _is_path_allowed(self, path: Path) -> bool:
        """
        Vérifier si un chemin est dans la liste autorisée
        Inclut protection contre symlinks et path traversal
        
        Security: Similar implementation to FileReaderTool for consistency
        
        Args:
            path: Path object to validate
            
        Returns:
            bool: True if path is allowed, False otherwise
        """
        try:
            path_resolved = path.resolve(strict=True)  # strict=True vérifie l'existence
        except (OSError, RuntimeError):
            return False

        # Vérifier chaque chemin autorisé
        for allowed in self.allowed_paths:
            try:
                allowed_resolved = Path(allowed).resolve()

                # Vérifier si le chemin est strictement sous le chemin autorisé
                path_resolved.relative_to(allowed_resolved)

                # Protection supplémentaire: vérifier les symlinks
                if path.is_symlink():
                    # Si c'est un symlink, vérifier que la cible est aussi dans l'allowlist
                    link_target = path.readlink()
                    if link_target.is_absolute():
                        # Lien absolu - doit être dans l'allowlist
                        target_resolved = link_target.resolve()
                        try:
                            target_resolved.relative_to(allowed_resolved)
                        except ValueError:
                            return False  # Lien pointe hors de l'allowlist
                    else:
                        # Lien relatif - résoudre depuis le répertoire du symlink
                        target_resolved = (path.parent / link_target).resolve()
                        try:
                            target_resolved.relative_to(allowed_resolved)
                        except ValueError:
                            return False  # Lien pointe hors de l'allowlist

                return True
            except ValueError:
                continue
        return False

    def execute(self, arguments: Dict[str, Any]) -> ToolResult:
        """
        Execute l'analyse de document

        Args:
            arguments: Dict avec 'file_path' et optionnellement 'analysis_type'

        Returns:
            ToolResult avec les données extraites
        """
        # Validate arguments
        is_valid, error = self.validate_arguments(arguments)
        if not is_valid:
            return ToolResult(status=ToolStatus.ERROR, output="", error=error)

        file_path = arguments["file_path"]
        analysis_type = arguments.get("analysis_type", "invoice")

        try:
            if analysis_type == "invoice":
                result = self.analyze_invoice(file_path)
            elif analysis_type == "financial":
                result = self.analyze_financial(file_path)
            elif analysis_type == "contract":
                result = self.analyze_contract(file_path)
            elif analysis_type == "report":
                result = self.analyze_report(file_path)
            else:  # extract or default
                result = self._extract_data(file_path)

            return ToolResult(status=ToolStatus.SUCCESS, output=str(result), metadata=result)
        except FileNotFoundError:
            # Redact PII from file path for compliance (Loi 25, PIPEDA)
            safe_path = self._redact_file_path(file_path)

            # Log at ERROR level for audit trail (Loi 25 Article 63.2)
            # Standard Python logging for testing compatibility
            logger.error(
                f"File not found error: {safe_path}",
                extra={
                    "actor": "document_analyzer_pme",
                    "event": "file.not_found",
                    "safe_path": safe_path,
                    "analysis_type": analysis_type,
                },
            )

            # Custom middleware logger for compliance
            if self.logger:
                self.logger.log_event(
                    actor="document_analyzer_pme",
                    event="file.not_found",
                    level="ERROR",
                    metadata={
                        "safe_path": safe_path,
                        "analysis_type": analysis_type,
                        "error": "FileNotFoundError",
                    },
                )

            return ToolResult(
                status=ToolStatus.ERROR, output="", error=f"File not found: {safe_path}"
            )
        except Exception as e:
            # Redact PII from error message
            safe_error_msg = self._redact_pii_from_error(str(e))
            safe_file_path = self._redact_file_path(file_path)

            # Log at ERROR level for audit trail (Loi 25 Article 63.2)
            # Standard Python logging for testing compatibility
            logger.error(
                f"Analysis error: {safe_error_msg}",
                extra={
                    "actor": "document_analyzer_pme",
                    "event": "analysis.error",
                    "safe_path": safe_file_path,
                    "error_type": type(e).__name__,
                    "analysis_type": analysis_type,
                },
            )

            # Custom middleware logger for compliance
            if self.logger:
                self.logger.log_event(
                    actor="document_analyzer_pme",
                    event="analysis.error",
                    level="ERROR",
                    metadata={
                        "safe_path": safe_file_path,
                        "error_type": type(e).__name__,
                        "error_message": safe_error_msg,
                        "analysis_type": analysis_type,
                    },
                )

            return ToolResult(
                status=ToolStatus.ERROR,
                output="",
                error=f"Error analyzing document: {safe_error_msg}",
            )

    def validate_arguments(self, arguments: Dict[str, Any]) -> tuple[bool, Optional[str]]:
        """Validate arguments with security checks"""
        if "file_path" not in arguments:
            return False, "Missing required argument: file_path"

        file_path = arguments["file_path"]
        if not isinstance(file_path, str):
            return False, "file_path must be a string"

        # Security: Check for null byte injection
        if '\0' in file_path:
            return False, "Null byte detected in path"

        # Security: Check path length (prevent buffer overflow)
        if len(file_path) > 4096:
            return False, "file_path too long (max 4096 characters)"

        # Security: Validate path is in allowlist (path traversal protection)
        # Check this BEFORE extension validation for better security
        try:
            path = Path(file_path)
            
            # Check if file exists before path validation
            # This prevents information leakage about file existence outside allowed paths
            if path.exists():
                if not self._is_path_allowed(path):
                    # Log blocked attempt for audit trail
                    if self.logger:
                        self.logger.log_event(
                            actor="document_analyzer_pme",
                            event="path.blocked",
                            level="WARNING",
                            metadata={
                                "attempted_path": self._redact_file_path(file_path),
                                "reason": "Path not in allowlist"
                            }
                        )
                    return False, f"Access denied: Path not in allowed directories"
        except Exception as e:
            return False, f"Path validation error: {str(e)}"

        # Check file extension
        valid_extensions = [".pdf", ".xlsx", ".xls", ".docx", ".doc"]
        if not any(file_path.lower().endswith(ext) for ext in valid_extensions):
            return False, f"Unsupported file type. Supported: {', '.join(valid_extensions)}"

        return True, None

    def _get_parameters_schema(self) -> Dict[str, Any]:
        """Return parameter schema"""
        return {
            "type": "object",
            "properties": {
                "file_path": {"type": "string", "description": "Path to the document to analyze"},
                "analysis_type": {
                    "type": "string",
                    "enum": ["invoice", "extract", "financial", "contract", "report"],
                    "description": "Type of analysis: invoice (TPS/TVQ), extract (raw data), financial (balance sheets/budgets), contract (legal clauses), report (general report)",
                },
            },
            "required": ["file_path"],
        }

    def analyze_invoice(self, file_path: str) -> Dict[str, Any]:
        """Analyse facture avec calculs taxes québécoises"""
        # Extraction données
        data = self._extract_data(file_path)

        # Calculs taxes
        subtotal = data.get("subtotal", 0)
        tps = subtotal * self.tps_rate
        tvq = subtotal * self.tvq_rate
        total = subtotal + tps + tvq

        return {
            "subtotal": subtotal,
            "tps": round(tps, 2),
            "tvq": round(tvq, 2),
            "total": round(total, 2),
            "tps_number": self._extract_tax_number(data, "TPS"),
            "tvq_number": self._extract_tax_number(data, "TVQ"),
            "pii_redacted": True,  # Conformité Loi 25
        }

    def analyze_financial(self, file_path: str) -> Dict[str, Any]:
        """
        Analyse financière pour bilans, budgets et états financiers

        Extrait les données financières clés et calcule des ratios pertinents
        pour les PME québécoises.
        """
        data = self._extract_data(file_path)
        text = data.get("text", "")

        # Extraction des montants financiers
        amounts = data.get("raw_amounts", [])
        numeric_amounts = []
        for amt in amounts:
            try:
                numeric_amounts.append(float(str(amt).replace(",", "")))
            except (ValueError, TypeError):
                continue

        # Calcul statistiques de base
        total_amounts = sum(numeric_amounts) if numeric_amounts else 0
        avg_amount = total_amounts / len(numeric_amounts) if numeric_amounts else 0
        max_amount = max(numeric_amounts) if numeric_amounts else 0
        min_amount = min(numeric_amounts) if numeric_amounts else 0

        # Detection de mots-cles financiers
        financial_keywords = {
            "actif": text.lower().count("actif"),
            "passif": text.lower().count("passif"),
            "capital": text.lower().count("capital"),
            "revenu": text.lower().count("revenu"),
            "depense": text.lower().count("dépense") + text.lower().count("depense"),
            "benefice": text.lower().count("bénéfice") + text.lower().count("benefice"),
            "perte": text.lower().count("perte"),
            "budget": text.lower().count("budget"),
        }

        return {
            "analysis_type": "financial",
            "document_type": "financial_statement",
            "amounts_detected": len(numeric_amounts),
            "total_amounts": round(total_amounts, 2),
            "average_amount": round(avg_amount, 2),
            "max_amount": round(max_amount, 2),
            "min_amount": round(min_amount, 2),
            "financial_keywords": financial_keywords,
            "data_source": data.get("columns", []) if "columns" in data else "text_extraction",
            "rows_analyzed": data.get("rows", 0) if "rows" in data else len(text.split("\n")),
            "pii_redacted": True,
        }

    def analyze_contract(self, file_path: str) -> Dict[str, Any]:
        """
        Analyse juridique de contrats

        Detecte les clauses importantes, dates, parties et obligations.
        Conforme aux exigences de la Loi 25 pour la protection des
        renseignements personnels.
        """
        data = self._extract_data(file_path)
        text = data.get("text", "")
        text_lower = text.lower()

        # Detection des parties contractantes (patterns rediges)
        parties_patterns = [
            r"entre\s+([^,]+),?\s+(?:ci-après|d'une part)",
            r"(?:le client|l'acheteur|le vendeur|le fournisseur)\s*:?\s*([^,\n]+)",
        ]
        parties_found = []
        for pattern in parties_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            parties_found.extend(matches[:2])  # Limiter a 2 par pattern

        # Detection de clauses importantes
        clause_keywords = {
            "confidentialite": "confidentialité" in text_lower or "confidentialite" in text_lower,
            "non_concurrence": "non-concurrence" in text_lower or "non concurrence" in text_lower,
            "resiliation": "résiliation" in text_lower or "resiliation" in text_lower,
            "garantie": "garantie" in text_lower,
            "responsabilite": "responsabilité" in text_lower or "responsabilite" in text_lower,
            "force_majeure": "force majeure" in text_lower,
            "propriete_intellectuelle": "propriété intellectuelle" in text_lower,
            "protection_donnees": "protection des données" in text_lower or "loi 25" in text_lower,
        }

        # Detection de dates
        date_patterns = [
            r"\d{1,2}[\s/-]\w+[\s/-]\d{4}",  # 15 janvier 2024
            r"\d{4}[\s/-]\d{2}[\s/-]\d{2}",  # 2024-01-15
            r"\d{1,2}/\d{1,2}/\d{4}",        # 15/01/2024
        ]
        dates_found = []
        for pattern in date_patterns:
            matches = re.findall(pattern, text)
            dates_found.extend(matches[:5])  # Limiter

        # Detection de montants
        amounts = data.get("raw_amounts", [])

        return {
            "analysis_type": "contract",
            "document_type": "legal_contract",
            "parties_detected": len(parties_found),
            "parties": ["[REDACTED]" for _ in parties_found],  # PII redaction
            "clauses_detected": clause_keywords,
            "important_clauses_count": sum(clause_keywords.values()),
            "dates_found": len(dates_found),
            "amounts_detected": len(amounts),
            "has_confidentiality": clause_keywords.get("confidentialite", False),
            "has_data_protection": clause_keywords.get("protection_donnees", False),
            "word_count": len(text.split()),
            "paragraphs": data.get("paragraphs", len(text.split("\n\n"))),
            "pii_redacted": True,
            "loi25_compliance_check": clause_keywords.get("protection_donnees", False),
        }

    def analyze_report(self, file_path: str) -> Dict[str, Any]:
        """
        Analyse generale de rapports

        Extrait la structure, les sections principales et les metriques
        cles du document.
        """
        data = self._extract_data(file_path)
        text = data.get("text", "")

        # Analyse de structure
        lines = text.split("\n")
        non_empty_lines = [l for l in lines if l.strip()]

        # Detection de sections/titres (lignes courtes en majuscules ou avec numerotation)
        sections = []
        for line in non_empty_lines:
            line_stripped = line.strip()
            # Titres potentiels: lignes courtes, numerotees, ou en majuscules
            if len(line_stripped) < 80:
                if (line_stripped.isupper() or
                    re.match(r"^\d+[\.\)]\s+", line_stripped) or
                    re.match(r"^[IVX]+[\.\)]\s+", line_stripped) or
                    line_stripped.startswith("#")):
                    sections.append(line_stripped[:50])  # Tronquer

        # Extraction de mots-cles de rapport
        report_keywords = {
            "introduction": "introduction" in text.lower(),
            "conclusion": "conclusion" in text.lower(),
            "recommandations": "recommandation" in text.lower(),
            "resume": "résumé" in text.lower() or "resume" in text.lower(),
            "analyse": "analyse" in text.lower(),
            "resultats": "résultat" in text.lower() or "resultat" in text.lower(),
            "methodologie": "méthodologie" in text.lower() or "methodologie" in text.lower(),
        }

        # Statistiques du document
        word_count = len(text.split())
        char_count = len(text)
        amounts = data.get("raw_amounts", [])

        return {
            "analysis_type": "report",
            "document_type": "general_report",
            "word_count": word_count,
            "character_count": char_count,
            "line_count": len(lines),
            "non_empty_lines": len(non_empty_lines),
            "sections_detected": len(sections),
            "section_titles": sections[:10],  # Limiter a 10 sections
            "structure_keywords": report_keywords,
            "has_standard_structure": sum(report_keywords.values()) >= 2,
            "amounts_detected": len(amounts),
            "estimated_pages": max(1, word_count // 300),  # ~300 mots/page
            "data_source": "excel" if "columns" in data else "text_document",
            "pii_redacted": True,
        }

    def _extract_data(self, file_path: str) -> Dict:
        """Extraction sécurisée avec redaction PII"""
        # Check file exists
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        # Security: Double-check path is allowed (defense in depth)
        if not self._is_path_allowed(path):
            raise PermissionError(f"Access denied: Path not in allowed directories")

        # Logique extraction selon type fichier
        if file_path.lower().endswith(".pdf"):
            return self._extract_pdf(file_path)
        elif file_path.lower().endswith((".xlsx", ".xls")):
            return self._extract_excel(file_path)
        elif file_path.lower().endswith((".docx", ".doc")):
            return self._extract_word(file_path)
        return {}

    def _extract_pdf(self, file_path: str) -> Dict:
        """Extract data from PDF file"""
        try:
            with open(file_path, "rb") as file:
                reader = PyPDF2.PdfReader(file)

                # Extract text from all pages
                text = ""
                for page in reader.pages:
                    text += page.extract_text()

                # Extract monetary values
                amounts = re.findall(r"\$?\s*(\d+(?:,\d{3})*(?:\.\d{2})?)", text)
                subtotal = 0
                if amounts:
                    # Try to find subtotal (usually largest or last before total)
                    try:
                        subtotal = float(amounts[-1].replace(",", ""))
                    except (ValueError, IndexError):
                        pass

                return {"text": text, "subtotal": subtotal, "raw_amounts": amounts}
        except Exception as e:
            raise ValueError(f"Error reading PDF: {str(e)}")

    def _extract_excel(self, file_path: str) -> Dict:  # noqa: C901
        """Extract data from Excel file"""
        try:
            df = pd.read_excel(file_path)

            # Look for common invoice columns
            subtotal = 0

            # Try to find subtotal in various ways
            for col in df.columns:
                col_lower = str(col).lower()
                if "subtotal" in col_lower or "sous-total" in col_lower:
                    values = df[col].dropna()
                    if not values.empty:
                        try:
                            subtotal = float(values.iloc[-1])
                        except (ValueError, TypeError):
                            pass
                    break

            # If no subtotal column, try to find numeric values
            if subtotal == 0:
                numeric_cols = df.select_dtypes(include=["float64", "int64"]).columns
                if len(numeric_cols) > 0:
                    values = df[numeric_cols[0]].dropna()
                    if not values.empty:
                        try:
                            subtotal = float(values.iloc[-1])
                        except (ValueError, TypeError):
                            pass

            return {
                "data": df.to_dict(),
                "subtotal": subtotal,
                "columns": list(df.columns),
                "rows": len(df),
            }
        except Exception as e:
            raise ValueError(f"Error reading Excel: {str(e)}")

    def _extract_word(self, file_path: str) -> Dict:
        """Extract data from Word document"""
        try:
            doc = docx.Document(file_path)

            # Extract all text
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])

            # Extract monetary values
            amounts = re.findall(r"\$?\s*(\d+(?:,\d{3})*(?:\.\d{2})?)", text)
            subtotal = 0
            if amounts:
                try:
                    subtotal = float(amounts[-1].replace(",", ""))
                except (ValueError, IndexError):
                    pass

            return {
                "text": text,
                "subtotal": subtotal,
                "paragraphs": len(doc.paragraphs),
                "raw_amounts": amounts,
            }
        except Exception as e:
            raise ValueError(f"Error reading Word document: {str(e)}")

    def _extract_tax_number(self, data: Dict, tax_type: str) -> str:
        """Extraction numéros taxes (TPS/TVQ)"""
        patterns = {"TPS": r"TPS[:\s]*(\d{9}RT\d{4})", "TVQ": r"TVQ[:\s]*(\d{10}TQ\d{4})"}

        # Search in text if available
        text = data.get("text", "")
        if text and tax_type in patterns:
            match = re.search(patterns[tax_type], text)
            if match:
                return "REDACTED"  # For security/compliance

        return "REDACTED"  # Par défaut pour sécurité
