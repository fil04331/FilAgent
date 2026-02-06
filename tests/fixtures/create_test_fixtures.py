"""
Script pour créer les fixtures de test pour le Document Analyzer
Phase 6.2.1: Création de fichiers de test pour tous les scénarios d'erreur
"""

import pandas as pd
from pathlib import Path
import io


def create_fixtures_directory():
    """Créer le répertoire fixtures s'il n'existe pas"""
    fixtures_dir = Path(__file__).parent
    fixtures_dir.mkdir(parents=True, exist_ok=True)
    return fixtures_dir


def create_valid_excel():
    """Créer un fichier Excel valide"""
    fixtures_dir = create_fixtures_directory()

    data = {
        "Description": ["Service A", "Service B", "Service C"],
        "Quantité": [10, 20, 5],
        "Prix unitaire": [100, 200, 150],
        "Montant": [1000, 4000, 750],
    }

    df = pd.DataFrame(data)
    output_path = fixtures_dir / "valid_invoice.xlsx"
    df.to_excel(output_path, index=False)
    print(f"✅ Créé: {output_path}")
    return output_path


def create_empty_excel():
    """Créer un fichier Excel vide"""
    fixtures_dir = create_fixtures_directory()

    df = pd.DataFrame()  # DataFrame vide
    output_path = fixtures_dir / "empty_file.xlsx"
    df.to_excel(output_path, index=False)
    print(f"✅ Créé: {output_path}")
    return output_path


def create_corrupted_excel():
    """Créer un fichier Excel corrompu"""
    fixtures_dir = create_fixtures_directory()

    output_path = fixtures_dir / "corrupted_file.xlsx"

    # Écrire du contenu invalide (pas un vrai fichier Excel)
    with open(output_path, "wb") as f:
        f.write(b"This is not a valid Excel file, just random text to simulate corruption.")

    print(f"✅ Créé: {output_path} (corrompu)")
    return output_path


def create_large_excel():
    """Créer un fichier Excel volumineux (> 50 MB)"""
    fixtures_dir = create_fixtures_directory()

    # Créer un DataFrame avec beaucoup de données
    # On vise environ 60 MB
    num_rows = 500000  # 500k lignes

    data = {
        "Column1": range(num_rows),
        "Column2": ["Data" * 100] * num_rows,  # Texte répété pour augmenter la taille
        "Column3": [i * 1.5 for i in range(num_rows)],
        "Column4": [f"Long text field {i} with additional content" * 10 for i in range(num_rows)],
    }

    df = pd.DataFrame(data)
    output_path = fixtures_dir / "large_file.xlsx"
    df.to_excel(output_path, index=False)

    size_mb = output_path.stat().st_size / (1024 * 1024)
    print(f"✅ Créé: {output_path} ({size_mb:.1f} MB)")
    return output_path


def create_valid_pdf():
    """Créer un fichier PDF valide (minimal)"""
    fixtures_dir = create_fixtures_directory()

    output_path = fixtures_dir / "valid_document.pdf"

    # Créer un PDF minimal valide manuellement
    # Structure PDF de base
    pdf_content = b"""%PDF-1.4
1 0 obj
<<
/Type /Catalog
/Pages 2 0 R
>>
endobj
2 0 obj
<<
/Type /Pages
/Kids [3 0 R]
/Count 1
>>
endobj
3 0 obj
<<
/Type /Page
/Parent 2 0 R
/Resources <<
/Font <<
/F1 <<
/Type /Font
/Subtype /Type1
/BaseFont /Helvetica
>>
>>
>>
/MediaBox [0 0 612 792]
/Contents 4 0 R
>>
endobj
4 0 obj
<<
/Length 150
>>
stream
BT
/F1 12 Tf
100 750 Td
(Facture Test) Tj
0 -20 Td
(Description: Services de consultation) Tj
0 -20 Td
(Montant: 1000.00 $) Tj
0 -20 Td
(TPS: 50.00 $) Tj
0 -20 Td
(Total: 1149.75 $) Tj
ET
endstream
endobj
xref
0 5
0000000000 65535 f
0000000009 00000 n
0000000058 00000 n
0000000115 00000 n
0000000317 00000 n
trailer
<<
/Size 5
/Root 1 0 R
>>
startxref
518
%%EOF
"""

    with open(output_path, "wb") as f:
        f.write(pdf_content)

    print(f"✅ Créé: {output_path}")
    return output_path


def create_corrupted_pdf():
    """Créer un fichier PDF corrompu"""
    fixtures_dir = create_fixtures_directory()

    output_path = fixtures_dir / "corrupted_document.pdf"

    # Écrire du contenu invalide (pas un vrai PDF)
    with open(output_path, "wb") as f:
        f.write(b"Not a valid PDF file - corrupted content here")

    print(f"✅ Créé: {output_path} (corrompu)")
    return output_path


def create_empty_pdf():
    """Créer un fichier PDF vide (0 bytes)"""
    fixtures_dir = create_fixtures_directory()

    output_path = fixtures_dir / "empty_document.pdf"

    # Créer un fichier vide
    output_path.touch()

    print(f"✅ Créé: {output_path} (vide)")
    return output_path


def create_valid_word():
    """Créer un fichier Word valide"""
    try:
        import docx

        fixtures_dir = create_fixtures_directory()

        doc = docx.Document()
        doc.add_heading("Rapport de Test", 0)
        doc.add_paragraph("Ceci est un document de test.")
        doc.add_paragraph("Il contient plusieurs paragraphes.")
        doc.add_heading("Section 1", level=1)
        doc.add_paragraph("Contenu de la section 1.")

        output_path = fixtures_dir / "valid_report.docx"
        doc.save(str(output_path))

        print(f"✅ Créé: {output_path}")
        return output_path
    except ImportError:
        print("⚠️ Module python-docx non disponible, skip création Word")
        return None


def create_empty_word():
    """Créer un fichier Word vide"""
    try:
        import docx

        fixtures_dir = create_fixtures_directory()

        doc = docx.Document()  # Document vide sans paragraphes

        output_path = fixtures_dir / "empty_report.docx"
        doc.save(str(output_path))

        print(f"✅ Créé: {output_path} (vide)")
        return output_path
    except ImportError:
        print("⚠️ Module python-docx non disponible, skip création Word vide")
        return None


def create_corrupted_word():
    """Créer un fichier Word corrompu"""
    fixtures_dir = create_fixtures_directory()

    output_path = fixtures_dir / "corrupted_report.docx"

    # Écrire du contenu invalide (pas un vrai fichier Word/ZIP)
    with open(output_path, "wb") as f:
        f.write(b"This is not a valid Word document - corrupted ZIP format")

    print(f"✅ Créé: {output_path} (corrompu)")
    return output_path


def create_unsupported_file():
    """Créer un fichier avec extension non supportée"""
    fixtures_dir = create_fixtures_directory()

    output_path = fixtures_dir / "unsupported_file.txt"

    with open(output_path, "w") as f:
        f.write("Ceci est un fichier texte, format non supporté par l'analyseur")

    print(f"✅ Créé: {output_path} (extension non supportée)")
    return output_path


def create_all_fixtures():
    """Créer toutes les fixtures de test"""
    print("=" * 60)
    print("🔧 Création des fixtures de test pour Document Analyzer")
    print("=" * 60)
    print()

    fixtures = {}

    # Excel
    print("📊 Fichiers Excel:")
    fixtures["valid_excel"] = create_valid_excel()
    fixtures["empty_excel"] = create_empty_excel()
    fixtures["corrupted_excel"] = create_corrupted_excel()
    print()

    # PDF
    print("📄 Fichiers PDF:")
    fixtures["valid_pdf"] = create_valid_pdf()
    fixtures["corrupted_pdf"] = create_corrupted_pdf()
    fixtures["empty_pdf"] = create_empty_pdf()
    print()

    # Word
    print("📝 Fichiers Word:")
    fixtures["valid_word"] = create_valid_word()
    fixtures["empty_word"] = create_empty_word()
    fixtures["corrupted_word"] = create_corrupted_word()
    print()

    # Autres
    print("📁 Autres fichiers:")
    fixtures["unsupported_file"] = create_unsupported_file()
    print()

    # Fichier volumineux (optionnel, peut prendre du temps)
    print("⚠️ Création fichier volumineux (peut prendre 30-60 secondes)...")
    # fixtures['large_excel'] = create_large_excel()  # Commenté par défaut
    print("ℹ️ Fichier volumineux skip (décommenter pour créer)")
    print()

    print("=" * 60)
    print(f"✅ {len([f for f in fixtures.values() if f is not None])} fixtures créées avec succès!")
    print("=" * 60)

    # Afficher la liste des fichiers créés
    print("\n📋 Liste des fixtures:")
    for name, path in fixtures.items():
        if path:
            size = path.stat().st_size if path.exists() else 0
            print(f"  - {name}: {path.name} ({size} bytes)")

    return fixtures


if __name__ == "__main__":
    create_all_fixtures()
