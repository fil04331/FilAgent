"""
Tests pour l'outil DocumentAnalyzerPME
Tests couvrant:
- Document parsing (PDF, Excel, Word)
- Content extraction
- Error handling for malformed documents
"""

import pytest
import sys
from pathlib import Path
import tempfile
import os

# Import PyPDF2 for creating test PDFs
try:
    import PyPDF2
    from PyPDF2 import PdfWriter
except ImportError:
    PyPDF2 = None
    PdfWriter = None

# Import pandas for creating test Excel files
try:
    import pandas as pd
except ImportError:
    pd = None

# Import docx for creating test Word files
try:
    import docx
except ImportError:
    docx = None

sys.path.insert(0, str(Path(__file__).parent.parent))

from tools.document_analyzer_pme import DocumentAnalyzerPME
from tools.base import ToolResult, ToolStatus

# ============================================================================
# FIXTURES
# ============================================================================


@pytest.fixture
def analyzer():
    """Create DocumentAnalyzerPME instance"""
    return DocumentAnalyzerPME()


@pytest.fixture
def temp_dir():
    """Create temporary directory for test files"""
    with tempfile.TemporaryDirectory() as tmpdir:
        yield tmpdir


@pytest.fixture
def sample_pdf(temp_dir):
    """Create a sample PDF file with invoice content"""
    if PyPDF2 is None:
        pytest.skip("PyPDF2 not available")

    pdf_path = os.path.join(temp_dir, "test_invoice.pdf")

    try:
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter

        # Create PDF with reportlab
        c = canvas.Canvas(pdf_path, pagesize=letter)
        c.drawString(100, 750, "FACTURE / INVOICE")
        c.drawString(100, 700, "Description: Services de consultation")
        c.drawString(100, 680, "Subtotal: $1000.00")
        c.drawString(100, 660, "TPS: 123456789RT0001")
        c.drawString(100, 640, "TVQ: 1234567890TQ0001")
        c.save()

        return pdf_path
    except ImportError:
        # Fallback: create a simple PDF with PyPDF2
        writer = PdfWriter()
        writer.add_blank_page(width=612, height=792)

        with open(pdf_path, "wb") as f:
            writer.write(f)

        return pdf_path


@pytest.fixture
def sample_excel(temp_dir):
    """Create a sample Excel file with invoice data"""
    if pd is None:
        pytest.skip("pandas not available")

    excel_path = os.path.join(temp_dir, "test_invoice.xlsx")

    data = {
        "Description": ["Item 1", "Item 2", "Item 3"],
        "Quantity": [1, 2, 3],
        "Price": [100.00, 200.00, 300.00],
        "Subtotal": [100.00, 400.00, 900.00],
    }

    df = pd.DataFrame(data)
    df.to_excel(excel_path, index=False)

    return excel_path


@pytest.fixture
def sample_word(temp_dir):
    """Create a sample Word document with invoice content"""
    if docx is None:
        pytest.skip("python-docx not available")

    word_path = os.path.join(temp_dir, "test_invoice.docx")

    doc = docx.Document()
    doc.add_heading("FACTURE / INVOICE", 0)
    doc.add_paragraph("Description: Services de consultation")
    doc.add_paragraph("Subtotal: $1000.00")
    doc.add_paragraph("TPS: 123456789RT0001")
    doc.add_paragraph("TVQ: 1234567890TQ0001")
    doc.save(word_path)

    return word_path


@pytest.fixture
def malformed_pdf(temp_dir):
    """Create a malformed PDF file"""
    pdf_path = os.path.join(temp_dir, "malformed.pdf")

    # Write invalid PDF content
    with open(pdf_path, "w") as f:
        f.write("This is not a valid PDF file")

    return pdf_path


@pytest.fixture
def malformed_excel(temp_dir):
    """Create a malformed Excel file"""
    excel_path = os.path.join(temp_dir, "malformed.xlsx")

    # Write invalid Excel content
    with open(excel_path, "w") as f:
        f.write("This is not a valid Excel file")

    return excel_path


@pytest.fixture
def malformed_word(temp_dir):
    """Create a malformed Word file"""
    word_path = os.path.join(temp_dir, "malformed.docx")

    # Write invalid Word content
    with open(word_path, "w") as f:
        f.write("This is not a valid Word file")

    return word_path


# ============================================================================
# BASIC TESTS
# ============================================================================


def test_analyzer_initialization(analyzer):
    """Test analyzer initialization"""
    assert analyzer.name == "document_analyzer_pme"
    assert analyzer.tps_rate == 0.05
    assert analyzer.tvq_rate == 0.09975


def test_analyzer_schema(analyzer):
    """Test analyzer schema"""
    schema = analyzer.get_schema()

    assert schema["name"] == "document_analyzer_pme"
    assert "description" in schema
    assert "parameters" in schema
    assert schema["parameters"]["type"] == "object"
    assert "file_path" in schema["parameters"]["properties"]
    assert "analysis_type" in schema["parameters"]["properties"]


# ============================================================================
# ARGUMENT VALIDATION TESTS
# ============================================================================


def test_validate_arguments_missing_file_path(analyzer):
    """Test validation with missing file_path"""
    is_valid, error = analyzer.validate_arguments({})

    assert not is_valid
    assert "file_path" in error.lower()


def test_validate_arguments_invalid_type(analyzer):
    """Test validation with invalid file_path type"""
    is_valid, error = analyzer.validate_arguments({"file_path": 123})

    assert not is_valid
    assert "string" in error.lower()


def test_validate_arguments_unsupported_extension(analyzer):
    """Test validation with unsupported file extension"""
    is_valid, error = analyzer.validate_arguments({"file_path": "test.txt"})

    assert not is_valid
    assert "unsupported" in error.lower()


def test_validate_arguments_valid_pdf(analyzer):
    """Test validation with valid PDF path"""
    is_valid, error = analyzer.validate_arguments({"file_path": "test.pdf"})

    assert is_valid
    assert error is None


def test_validate_arguments_valid_excel(analyzer):
    """Test validation with valid Excel path"""
    is_valid, error = analyzer.validate_arguments({"file_path": "test.xlsx"})

    assert is_valid
    assert error is None


def test_validate_arguments_valid_word(analyzer):
    """Test validation with valid Word path"""
    is_valid, error = analyzer.validate_arguments({"file_path": "test.docx"})

    assert is_valid
    assert error is None


# ============================================================================
# DOCUMENT PARSING TESTS - PDF
# ============================================================================


@pytest.mark.skipif(PyPDF2 is None, reason="PyPDF2 not available")
def test_extract_pdf_basic(analyzer, sample_pdf):
    """Test basic PDF extraction"""
    data = analyzer._extract_pdf(sample_pdf)

    assert "text" in data
    assert "subtotal" in data
    assert isinstance(data["subtotal"], (int, float))


@pytest.mark.skipif(PyPDF2 is None, reason="PyPDF2 not available")
def test_extract_pdf_with_amounts(analyzer, sample_pdf):
    """Test PDF extraction with amount detection"""
    data = analyzer._extract_pdf(sample_pdf)

    assert "raw_amounts" in data
    # Should have found some amounts
    assert isinstance(data["raw_amounts"], list)


# ============================================================================
# DOCUMENT PARSING TESTS - EXCEL
# ============================================================================


@pytest.mark.skipif(pd is None, reason="pandas not available")
def test_extract_excel_basic(analyzer, sample_excel):
    """Test basic Excel extraction"""
    data = analyzer._extract_excel(sample_excel)

    assert "data" in data
    assert "subtotal" in data
    assert "columns" in data
    assert "rows" in data


@pytest.mark.skipif(pd is None, reason="pandas not available")
def test_extract_excel_subtotal(analyzer, sample_excel):
    """Test Excel subtotal extraction"""
    data = analyzer._extract_excel(sample_excel)

    # Should find the subtotal column
    assert data["subtotal"] > 0
    # Should be the last value in subtotal column (900.00)
    assert data["subtotal"] == 900.00


@pytest.mark.skipif(pd is None, reason="pandas not available")
def test_extract_excel_columns(analyzer, sample_excel):
    """Test Excel column extraction"""
    data = analyzer._extract_excel(sample_excel)

    assert "Description" in data["columns"]
    assert "Quantity" in data["columns"]
    assert "Price" in data["columns"]
    assert "Subtotal" in data["columns"]


@pytest.mark.skipif(pd is None, reason="pandas not available")
def test_extract_excel_rows(analyzer, sample_excel):
    """Test Excel row counting"""
    data = analyzer._extract_excel(sample_excel)

    assert data["rows"] == 3


# ============================================================================
# DOCUMENT PARSING TESTS - WORD
# ============================================================================


@pytest.mark.skipif(docx is None, reason="python-docx not available")
def test_extract_word_basic(analyzer, sample_word):
    """Test basic Word extraction"""
    data = analyzer._extract_word(sample_word)

    assert "text" in data
    assert "subtotal" in data
    assert "paragraphs" in data


@pytest.mark.skipif(docx is None, reason="python-docx not available")
def test_extract_word_text(analyzer, sample_word):
    """Test Word text extraction"""
    data = analyzer._extract_word(sample_word)

    assert "FACTURE" in data["text"] or "INVOICE" in data["text"]


@pytest.mark.skipif(docx is None, reason="python-docx not available")
def test_extract_word_amounts(analyzer, sample_word):
    """Test Word amount extraction"""
    data = analyzer._extract_word(sample_word)

    assert "raw_amounts" in data
    assert isinstance(data["raw_amounts"], list)


# ============================================================================
# INVOICE ANALYSIS TESTS
# ============================================================================


@pytest.mark.skipif(PyPDF2 is None, reason="PyPDF2 not available")
def test_analyze_invoice_pdf(analyzer, sample_pdf):
    """Test invoice analysis from PDF"""
    result = analyzer.analyze_invoice(sample_pdf)

    assert "subtotal" in result
    assert "tps" in result
    assert "tvq" in result
    assert "total" in result
    assert "pii_redacted" in result
    assert result["pii_redacted"] is True


def test_analyze_invoice_tax_calculations(analyzer):
    """Test tax calculation logic"""

    # Create mock data
    class MockAnalyzer(DocumentAnalyzerPME):
        def _extract_data(self, file_path):
            return {"subtotal": 1000.00, "text": ""}

    mock_analyzer = MockAnalyzer()
    result = mock_analyzer.analyze_invoice("dummy.pdf")

    # Check calculations
    assert result["subtotal"] == 1000.00
    assert result["tps"] == 50.00  # 5% of 1000
    assert result["tvq"] == 99.75  # 9.975% of 1000
    assert result["total"] == 1149.75  # 1000 + 50 + 99.75


def test_analyze_invoice_tax_numbers_redacted(analyzer):
    """Test that tax numbers are redacted for compliance"""

    class MockAnalyzer(DocumentAnalyzerPME):
        def _extract_data(self, file_path):
            return {"subtotal": 1000.00, "text": "TPS: 123456789RT0001 TVQ: 1234567890TQ0001"}

    mock_analyzer = MockAnalyzer()
    result = mock_analyzer.analyze_invoice("dummy.pdf")

    # Tax numbers should be redacted
    assert result["tps_number"] == "REDACTED"
    assert result["tvq_number"] == "REDACTED"


# ============================================================================
# EXECUTE METHOD TESTS
# ============================================================================


@pytest.mark.skipif(PyPDF2 is None, reason="PyPDF2 not available")
def test_execute_invoice_analysis(analyzer, sample_pdf):
    """Test execute method with invoice analysis"""
    result = analyzer.execute({"file_path": sample_pdf, "analysis_type": "invoice"})

    assert result.is_success()
    assert result.metadata is not None
    assert "subtotal" in result.metadata
    assert "tps" in result.metadata
    assert "tvq" in result.metadata


@pytest.mark.skipif(PyPDF2 is None, reason="PyPDF2 not available")
def test_execute_extract_mode(analyzer, sample_pdf):
    """Test execute method with extract mode"""
    result = analyzer.execute({"file_path": sample_pdf, "analysis_type": "extract"})

    assert result.is_success()
    assert result.metadata is not None


def test_execute_missing_file(analyzer):
    """Test execute with non-existent file"""
    result = analyzer.execute({"file_path": "/nonexistent/file.pdf"})

    assert not result.is_success()
    assert result.status == ToolStatus.ERROR
    assert "not found" in result.error.lower()


def test_execute_invalid_arguments(analyzer):
    """Test execute with invalid arguments"""
    result = analyzer.execute({})

    assert not result.is_success()
    assert result.status == ToolStatus.ERROR
    assert "file_path" in result.error.lower()


# ============================================================================
# ERROR HANDLING TESTS - MALFORMED DOCUMENTS
# ============================================================================


def test_malformed_pdf_handling(analyzer, malformed_pdf):
    """Test handling of malformed PDF file"""
    result = analyzer.execute({"file_path": malformed_pdf})

    assert not result.is_success()
    assert result.status == ToolStatus.ERROR
    assert "error" in result.error.lower()


def test_malformed_excel_handling(analyzer, malformed_excel):
    """Test handling of malformed Excel file"""
    result = analyzer.execute({"file_path": malformed_excel})

    assert not result.is_success()
    assert result.status == ToolStatus.ERROR


def test_malformed_word_handling(analyzer, malformed_word):
    """Test handling of malformed Word file"""
    result = analyzer.execute({"file_path": malformed_word})

    assert not result.is_success()
    assert result.status == ToolStatus.ERROR


def test_empty_pdf_handling(analyzer, temp_dir):
    """Test handling of empty PDF file"""
    if PyPDF2 is None:
        pytest.skip("PyPDF2 not available")

    empty_pdf = os.path.join(temp_dir, "empty.pdf")

    # Create minimal valid PDF
    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)

    with open(empty_pdf, "wb") as f:
        writer.write(f)

    result = analyzer.execute({"file_path": empty_pdf})

    # Should succeed but with zero/minimal data
    assert result.is_success()
    assert result.metadata is not None


# ============================================================================
# TAX NUMBER EXTRACTION TESTS
# ============================================================================


def test_extract_tax_number_tps_pattern(analyzer):
    """Test TPS number pattern extraction"""
    data = {"text": "TPS: 123456789RT0001 Some other text"}

    result = analyzer._extract_tax_number(data, "TPS")

    # Should be redacted for compliance
    assert result == "REDACTED"


def test_extract_tax_number_tvq_pattern(analyzer):
    """Test TVQ number pattern extraction"""
    data = {"text": "TVQ: 1234567890TQ0001 Some other text"}

    result = analyzer._extract_tax_number(data, "TVQ")

    # Should be redacted for compliance
    assert result == "REDACTED"


def test_extract_tax_number_no_match(analyzer):
    """Test tax number extraction with no match"""
    data = {"text": "No tax numbers here"}

    result = analyzer._extract_tax_number(data, "TPS")

    # Should return REDACTED by default
    assert result == "REDACTED"


def test_extract_tax_number_no_text(analyzer):
    """Test tax number extraction with no text"""
    data = {}

    result = analyzer._extract_tax_number(data, "TPS")

    # Should return REDACTED by default
    assert result == "REDACTED"


# ============================================================================
# EDGE CASES AND SPECIAL SCENARIOS
# ============================================================================


def test_case_insensitive_file_extension(analyzer):
    """Test that file extension checking is case-insensitive"""
    is_valid_lower, _ = analyzer.validate_arguments({"file_path": "test.pdf"})
    is_valid_upper, _ = analyzer.validate_arguments({"file_path": "test.PDF"})
    is_valid_mixed, _ = analyzer.validate_arguments({"file_path": "test.PdF"})

    assert is_valid_lower
    assert is_valid_upper
    assert is_valid_mixed


def test_analyze_invoice_zero_subtotal(analyzer):
    """Test invoice analysis with zero subtotal"""

    class MockAnalyzer(DocumentAnalyzerPME):
        def _extract_data(self, file_path):
            return {"subtotal": 0, "text": ""}

    mock_analyzer = MockAnalyzer()
    result = mock_analyzer.analyze_invoice("dummy.pdf")

    assert result["subtotal"] == 0
    assert result["tps"] == 0.0
    assert result["tvq"] == 0.0
    assert result["total"] == 0.0


@pytest.mark.skipif(pd is None, reason="pandas not available")
def test_extract_excel_no_subtotal_column(analyzer, temp_dir):
    """Test Excel extraction when no subtotal column exists"""
    excel_path = os.path.join(temp_dir, "no_subtotal.xlsx")

    data = {"Item": ["A", "B", "C"], "Amount": [100, 200, 300]}

    df = pd.DataFrame(data)
    df.to_excel(excel_path, index=False)

    result = analyzer._extract_excel(excel_path)

    # Should still extract data, using numeric column
    assert "data" in result
    assert "subtotal" in result


@pytest.mark.skipif(pd is None, reason="pandas not available")
def test_extract_excel_french_column_names(analyzer, temp_dir):
    """Test Excel extraction with French column names"""
    excel_path = os.path.join(temp_dir, "french.xlsx")

    data = {
        "Description": ["Item 1", "Item 2"],
        "Quantité": [1, 2],
        "Prix": [100, 200],
        "Sous-total": [100, 400],
    }

    df = pd.DataFrame(data)
    df.to_excel(excel_path, index=False)

    result = analyzer._extract_excel(excel_path)

    # Should recognize "Sous-total" as subtotal
    assert result["subtotal"] == 400


# ============================================================================
# COMPLIANCE TESTS
# ============================================================================


def test_pii_redaction_enabled(analyzer):
    """Test that PII redaction is always enabled"""

    class MockAnalyzer(DocumentAnalyzerPME):
        def _extract_data(self, file_path):
            return {"subtotal": 1000.00, "text": ""}

    mock_analyzer = MockAnalyzer()
    result = mock_analyzer.analyze_invoice("dummy.pdf")

    assert result["pii_redacted"] is True


def test_tax_numbers_always_redacted(analyzer):
    """Test that tax numbers are always redacted (Loi 25 compliance)"""

    class MockAnalyzer(DocumentAnalyzerPME):
        def _extract_data(self, file_path):
            return {"subtotal": 1000.00, "text": "TPS: 123456789RT0001 TVQ: 1234567890TQ0001"}

    mock_analyzer = MockAnalyzer()
    result = mock_analyzer.analyze_invoice("dummy.pdf")

    # Tax numbers must be redacted
    assert result["tps_number"] == "REDACTED"
    assert result["tvq_number"] == "REDACTED"
    # Ensure they're not in the output
    assert "123456789RT0001" not in str(result)
    assert "1234567890TQ0001" not in str(result)
